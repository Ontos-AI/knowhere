"""Generate a public, synthetic DOCX for standalone Codex export tests."""

from __future__ import annotations

from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


def _make_fixture_image(path: Path) -> None:
    document = pymupdf.open()
    try:
        page = document.new_page(width=240, height=140)
        page.draw_rect((4, 4, 236, 136), color=(0.1, 0.3, 0.8), fill=(0.8, 0.9, 1))
        page.draw_circle((120, 70), 38, color=(0.8, 0.2, 0.1), fill=(1, 0.7, 0.5))
        pixmap = page.get_pixmap(matrix=pymupdf.Matrix(1, 1), alpha=False)
        pixmap.save(path)
    finally:
        document.close()


def generate_docx_fixture(output_path: Path) -> Path:
    """Write a DOCX with hierarchy, tables, image, rich text, and page break."""
    destination = output_path.expanduser().resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    fixture_image = destination.with_name(f".{destination.stem}-fixture-image.png")
    _make_fixture_image(fixture_image)
    try:
        document = Document()
        document.core_properties.title = "Synthetic DOCX Technical Review"
        document.core_properties.author = "Knowhere test fixture"
        document.add_heading("1. Overview", level=1)
        document.add_paragraph(
            "This public synthetic fixture contains no client or confidential content."
        )

        rich = document.add_paragraph("Rich text: H")
        subscript = rich.add_run("2")
        subscript.font.subscript = True
        rich.add_run("O and 10")
        superscript = rich.add_run("-6")
        superscript.font.superscript = True
        rich.add_run(" A; symbols ≤ ± µA °C.")

        simple = document.add_table(rows=3, cols=2)
        simple.style = "Table Grid"
        values = (("Metric", "Value"), ("Current", "5 µA"), ("Temperature", "25 °C"))
        for row, values_row in zip(simple.rows, values, strict=True):
            for cell, value in zip(row.cells, values_row, strict=True):
                cell.text = value

        document.add_picture(str(fixture_image), width=Inches(2.5))
        document.paragraphs[-1].add_run(" Synthetic embedded image")
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        document.add_heading("1.1 Results", level=2)
        document.add_paragraph("Merged cells exercise complex table preservation.")
        complex_table = document.add_table(rows=3, cols=3)
        complex_table.style = "Table Grid"
        complex_table.cell(0, 0).merge(complex_table.cell(1, 0)).text = "Group"
        complex_table.cell(0, 1).merge(complex_table.cell(0, 2)).text = "Limits"
        complex_table.cell(1, 1).text = "Low"
        complex_table.cell(1, 2).text = "High"
        complex_table.cell(2, 0).text = "A"
        complex_table.cell(2, 1).text = "1"
        complex_table.cell(2, 2).text = "9"
        document.add_heading("2. Conclusion", level=1)
        document.add_paragraph(
            "Use extracted structure for navigation and verify evidence in the native DOCX."
        )
        footer = document.sections[0].footer.paragraphs[0]
        footer.text = "Synthetic fixture — no client content"
        document.save(destination)
    finally:
        fixture_image.unlink(missing_ok=True)
    return destination


if __name__ == "__main__":
    generate_docx_fixture(Path(__file__).with_name("synthetic-review.docx"))
